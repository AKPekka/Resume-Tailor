import streamlit as st
import tempfile
import os
import torch
from document_processor import DocumentProcessor
from keyword_extractor import KeywordExtractor
from resume_tailor import ResumeTailor
from sbert import SemanticMatcher
from gpt2 import ResumeBulletGenerator
import docx
from docx.shared import RGBColor

def main():
    st.set_page_config(
        page_title="AI-Enhanced Resume Tailoring Tool",
        page_icon="📝",
        layout="wide"
    )

    st.title("AKPs AI-Enhanced Resume Tailoring Tool")
    st.write("""
    This tool uses advanced AI models (BERT and GPT-2) to help you tailor your resume
    to match job descriptions. Upload your resume and a job description to get started.
    """)

    # Check for GPU availability
    device = "cuda" if torch.cuda.is_available() else "cpu"
    st.sidebar.info(f"Running on: {device.upper()}")

    # Model selection
    st.sidebar.header("Model Configuration")
    sbert_model_name = st.sidebar.selectbox(
        "Select SBERT Model",
        ["paraphrase-MiniLM-L6-v2", "all-mpnet-base-v2"],
        index=1, # Default to the better SBERT model
        help="'all-mpnet-base-v2' is generally more accurate but slower."
    )

    generator_model_name = st.sidebar.selectbox(
        "Select Bullet Generator Model",
        ["gpt2-medium", "deepseek-ai/deepseek-coder-1.3b-instruct"],
        index=1,  # Default to DeepSeek Coder
        help="'deepseek-ai/deepseek-coder-1.3b-instruct' is a more advanced model for generation."
    )

    matching_threshold = st.sidebar.slider(
        "Matching Threshold",
        min_value=0.5,
        max_value=0.9,
        value=0.7,
        help="Lower values find more potential matches, higher values are more strict"
    )

    # Initialize processors
    doc_processor = DocumentProcessor()
    keyword_extractor = KeywordExtractor()
    resume_tailor = ResumeTailor()

    # Initialize AI models with loading state indicators
    sbert_load_state = st.sidebar.text("SBERT: Not loaded")
    generator_load_state = st.sidebar.text("Generator: Not loaded") # Changed from gpt2_load_state

    # Create session state for storing data
    if 'tmp_resume_path' not in st.session_state:
        st.session_state.tmp_resume_path = None
    if 'tmp_job_path' not in st.session_state:
        st.session_state.tmp_job_path = None
    if 'resume_type' not in st.session_state:
        st.session_state.resume_type = None
    if 'resume_text' not in st.session_state:
        st.session_state.resume_text = None
    if 'job_text' not in st.session_state:
        st.session_state.job_text = None
    if 'keyword_results' not in st.session_state:
        st.session_state.keyword_results = None
    if 'semantic_results' not in st.session_state:
        st.session_state.semantic_results = None
    if 'suggestions' not in st.session_state:
        st.session_state.suggestions = None
    if 'gpt2_suggestions' not in st.session_state: # Will be 'generator_suggestions' effectively
        st.session_state.gpt2_suggestions = None 
    if 'models_loaded' not in st.session_state:
        st.session_state.models_loaded = False

    # File upload section
    st.header("1. Upload Documents")
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Upload Your Resume")
        resume_file = st.file_uploader("Choose your resume file", type=["pdf", "docx"])

    with col2:
        st.subheader("Upload Job Description")
        job_desc_file = st.file_uploader("Choose the job description file", type=["pdf", "docx"])

    # Process documents if uploaded
    if resume_file and job_desc_file:
        # First, load the AI models if not already loaded
        if not st.session_state.models_loaded:
            with st.spinner("Loading AI models..."):
                try:
                    # Initialize models
                    sbert_load_state.text(f"SBERT ({sbert_model_name}): Loading...")
                    st.session_state.semantic_matcher = SemanticMatcher(model_name=sbert_model_name)
                    sbert_load_state.text(f"SBERT ({sbert_model_name}): ✅ Loaded")

                    generator_load_state.text(f"Generator ({generator_model_name}): Loading...")
                    st.session_state.bullet_generator = ResumeBulletGenerator(
                        keyword_extractor_instance=keyword_extractor, 
                        model_name=generator_model_name
                    )
                    generator_load_state.text(f"Generator ({generator_model_name}): ✅ Loaded")

                    st.session_state.models_loaded = True
                except Exception as e:
                    st.error(f"Error loading AI models: {str(e)}")
                    st.exception(e) # Print full traceback to terminal
                    st.stop()

        print("Starting document analysis...") # DEBUG
        with st.spinner("Analyzing documents with AI..."):
            # Process resume
            print("Step 1: Processing resume file...") # DEBUG
            resume_type = 'pdf' if resume_file.name.endswith('.pdf') else 'docx'
            st.session_state.resume_type = resume_type

            # Create temporary files
            tmp_resume = tempfile.NamedTemporaryFile(delete=False, suffix=f".{resume_type}")
            tmp_resume.write(resume_file.getvalue())
            tmp_resume.close()
            st.session_state.tmp_resume_path = tmp_resume.name

            # Process job description
            print("Step 2: Processing job description file...") # DEBUG
            job_type = 'pdf' if job_desc_file.name.endswith('.pdf') else 'docx'
            tmp_job = tempfile.NamedTemporaryFile(delete=False, suffix=f".{job_type}")
            tmp_job.write(job_desc_file.getvalue())
            tmp_job.close()
            st.session_state.tmp_job_path = tmp_job.name

            # Extract text
            print("Step 3: Extracting text from resume...") # DEBUG
            resume_text = doc_processor.extract_text(st.session_state.tmp_resume_path, resume_type)
            st.session_state.resume_text = resume_text
            print(f"Resume text extracted (length: {len(resume_text)} chars)") # DEBUG

            print("Step 4: Extracting text from job description...") # DEBUG
            job_text = doc_processor.extract_text(st.session_state.tmp_job_path, job_type)
            st.session_state.job_text = job_text
            print(f"Job text extracted (length: {len(job_text)} chars)") # DEBUG

            # 1. Simple keyword extraction (original functionality)
            print("Step 5: Performing keyword extraction...") # DEBUG
            st.session_state.keyword_results = keyword_extractor.get_keyword_suggestions(
                resume_text, job_text, similarity_threshold=matching_threshold 
            )
            job_keywords_for_semantic = st.session_state.keyword_results.get('job_keywords', [])
            print(f"Keyword extraction complete. Found {len(job_keywords_for_semantic)} job keywords.") # DEBUG

            # 2. Semantic matching using SBERT
            print("Step 6: Performing semantic matching (SBERT)...") # DEBUG
            st.session_state.semantic_results = st.session_state.semantic_matcher.find_missing_skills(
                resume_text=resume_text, 
                job_text=job_text,
                job_keywords=job_keywords_for_semantic,
                similarity_threshold=matching_threshold 
            )
            print("Semantic matching complete.") # DEBUG
            if st.session_state.semantic_results:
                 print(f"Missing skills found: {len(st.session_state.semantic_results.get('missing_skills',[]))}, Weak skills: {len(st.session_state.semantic_results.get('weak_skills',[]))}")

            # 3. Generate tailored suggestions using ResumeTailor
            print("Step 7: Generating basic tailoring suggestions (ResumeTailor)...") # DEBUG
            st.session_state.suggestions = resume_tailor.generate_tailoring_suggestions(
                st.session_state.keyword_results, resume_text
            )
            print("Basic tailoring suggestions complete.") # DEBUG

            # 4. Generate resume bullet points using the selected generator
            print(f"Step 8: Generating bullet points with {generator_model_name}...") # DEBUG
            st.session_state.gpt2_suggestions = st.session_state.bullet_generator.generate_suggestions(
                st.session_state.semantic_results, job_text, resume_text
            )
            print("Bullet point generation complete.") # DEBUG
        
        print("Document analysis finished.") # DEBUG

        # Display results
        st.header("2. Analysis Results")

        # Create tabs for different analysis methods
        tab1, tab2, tab3 = st.tabs(["Keyword Analysis", "Semantic Analysis (BERT)", "AI-Generated Bullet Points"])

        with tab1:
            # Display similarity score
            keyword_similarity = st.session_state.keyword_results['similarity_score']
            st.subheader("Keyword-Based Similarity")

            # Create a color for the similarity score
            if keyword_similarity < 0.5:
                color = "red"
            elif keyword_similarity < 0.7:
                color = "orange"
            else:
                color = "green"

            st.markdown(f"<h3 style='color: {color};'>{keyword_similarity:.2f}/1.00</h3>", unsafe_allow_html=True)

            # Display keywords
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("Keywords in Your Resume")
                for keyword in st.session_state.keyword_results['resume_keywords']:
                    st.write(f"- {keyword}")

            with col2:
                st.subheader("Keywords in Job Description")
                for keyword in st.session_state.keyword_results['job_keywords']:
                    if keyword in st.session_state.keyword_results['missing_keywords']:
                        st.markdown(f"- <span style='color: red;'>{keyword} (missing)</span>", unsafe_allow_html=True)
                    else:
                        st.write(f"- {keyword}")

        with tab2:
            # Display BERT semantic matching results
            semantic_similarity = st.session_state.semantic_results['similarity_score']
            st.subheader("BERT Semantic Similarity")

            # Create a color for the semantic similarity score
            if semantic_similarity < 0.5:
                color = "red"
            elif semantic_similarity < 0.7:
                color = "orange"
            else:
                color = "green"

            st.markdown(f"<h3 style='color: {color};'>{semantic_similarity:.2f}/1.00</h3>", unsafe_allow_html=True)

            # Display semantically matched skills
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("Missing Skills (Semantic)")
                if st.session_state.semantic_results['missing_skills']:
                    for skill in st.session_state.semantic_results['missing_skills']:
                        confidence = skill['confidence'] * 100
                        st.markdown(f"- **{skill['skill']}** (Confidence: {confidence:.1f}%)")
                else:
                    st.write("No missing skills detected!")

            with col2:
                st.subheader("Weak Skills (Need Improvement)")
                if st.session_state.semantic_results['weak_skills']:
                    for skill in st.session_state.semantic_results['weak_skills']:
                        confidence = skill['confidence'] * 100
                        st.markdown(f"- **{skill['skill']}** (Confidence: {confidence:.1f}%)")
                        with st.expander("Current Context"):
                            st.write(skill['context'])
                else:
                    st.write("No weak skills detected!")

        with tab3:
            st.subheader("Resume Bullet Point Suggestions (AI Generated)")
            if st.session_state.gpt2_suggestions:
                # Display suggestions for missing skills
                if st.session_state.gpt2_suggestions.get('missing_skills'):
                    st.markdown("#### For Missing Skills:")
                    for item in st.session_state.gpt2_suggestions['missing_skills']:
                        st.markdown(f"**Skill:** {item['skill']}")
                        if item['bullet_points']:
                            for bullet in item['bullet_points']:
                                st.markdown(f"- {bullet}")
                        else:
                            st.markdown("_(No specific bullet points generated for this skill)_ ")
                        st.markdown("---")
                else:
                    st.markdown("_No missing skills found to generate bullet points for._")

                # Display suggestions for weak skills
                if st.session_state.gpt2_suggestions.get('weak_skills'):
                    st.markdown("#### For Weak Skills (Improvements):")
                    for item in st.session_state.gpt2_suggestions['weak_skills']:
                        st.markdown(f"**Skill:** {item['skill']}")
                        st.markdown(f"**Original Context Snippet:** {item.get('current_context', 'N/A')}")
                        if item['improved_bullet_points']:
                            st.markdown("**Suggested Improvements:**")
                            for bullet in item['improved_bullet_points']:
                                st.markdown(f"- {bullet}")
                        else:
                            st.markdown("_(No specific improvement suggestions generated for this skill)_ ")
                        st.markdown("---")
                else:
                    st.markdown("_No weak skills found to generate improvements for._")
            else:
                st.info("Upload documents to see AI-generated bullet point suggestions.")

        # Create tailored resume
        st.header("4. Create Tailored Resume")

        if st.button("Generate AI-Enhanced Resume"):
            with st.spinner("Creating tailored resume with AI suggestions..."):
                # Only works with DOCX files
                if st.session_state.resume_type == 'docx':
                    try:
                        # Load the original resume
                        doc = docx.Document(st.session_state.tmp_resume_path)

                        # Add a section for AI suggestions
                        doc.add_paragraph().add_run("AI-GENERATED SUGGESTIONS:").bold = True
                        doc.add_paragraph("The following are AI-generated suggestions to improve your resume:")

                        # Add missing skills with GPT-2 bullet points
                        p = doc.add_paragraph()
                        run = p.add_run("MISSING SKILLS - Consider adding these bullet points:")
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red

                        for skill_suggestion in st.session_state.gpt2_suggestions['missing_skills']:
                            skill = skill_suggestion['skill']
                            doc.add_paragraph(f"{skill.title()}:").bold = True

                            for bullet in skill_suggestion['bullet_points']:
                                doc.add_paragraph(f"• {bullet}", style='ListBullet')

                        # Add weak skills with improved bullet points
                        p = doc.add_paragraph()
                        run = p.add_run("WEAK SKILLS - Consider replacing with these improved versions:")
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 128, 0)  # Orange

                        for skill_suggestion in st.session_state.gpt2_suggestions['weak_skills']:
                            skill = skill_suggestion['skill']
                            doc.add_paragraph(f"{skill.title()}:").bold = True

                            for bullet in skill_suggestion['improved_bullet_points']:
                                doc.add_paragraph(f"• {bullet}", style='ListBullet')

                        # Add original suggestions
                        p = doc.add_paragraph()
                        run = p.add_run("ADDITIONAL SUGGESTIONS:")
                        run.bold = True

                        for suggestion in st.session_state.suggestions:
                            if suggestion['keyword']:
                                p = doc.add_paragraph(f"• {suggestion['suggestion']} (Missing keyword: '{suggestion['keyword']}')")
                            else:
                                p = doc.add_paragraph(f"• {suggestion['suggestion']}")

                        # Save the modified document
                        tmp_modified = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        tmp_modified_path = tmp_modified.name
                        tmp_modified.close()

                        doc.save(tmp_modified_path)

                        # Provide download link
                        with open(tmp_modified_path, "rb") as file:
                            st.download_button(
                                label="Download AI-Enhanced Resume",
                                data=file,
                                file_name="ai_enhanced_resume.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                        # Clean up modified file (keep original files for potential reuse)
                        os.unlink(tmp_modified_path)
                    except Exception as e:
                        st.error(f"Error creating tailored resume: {str(e)}")
                else:
                    st.error("Tailoring feature currently only works with DOCX files. Please upload a DOCX resume.")

    # Clean up temporary files when the session ends
    if st.button("Clear All Data"):
        if st.session_state.tmp_resume_path and os.path.exists(st.session_state.tmp_resume_path):
            try:
                os.unlink(st.session_state.tmp_resume_path)
            except:
                pass

        if st.session_state.tmp_job_path and os.path.exists(st.session_state.tmp_job_path):
            try:
                os.unlink(st.session_state.tmp_job_path)
            except:
                pass

        # Reset session state
        for key in list(st.session_state.keys()):
            del st.session_state[key]

        st.experimental_rerun()

if __name__ == "__main__":
    main()
